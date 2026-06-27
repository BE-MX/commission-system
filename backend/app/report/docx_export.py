"""生产订单 Word 导出服务。

使用 python-docx 生成 .docx 文件，支持 A4/A3/A5/B5 + 横版/竖版。
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, Emu, RGBColor

# 模板资源目录（与 HTML 打印模板共享）
_ASSETS_DIR = Path(__file__).parent / "templates" / "assets"
_MIDDLE_PUNCH_IMAGE = _ASSETS_DIR / "middle_punch_requirement.jpg"
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 纸张尺寸映射 (宽 cm, 高 cm)
PAGE_SIZES = {
    "A4": (21.0, 29.7),
    "A3": (29.7, 42.0),
    "A5": (14.8, 21.0),
    "B5": (17.6, 25.0),
}

# 字体 & 颜色
_FONT = "Microsoft YaHei"
_COLOR_ACCENT = RGBColor(0xD4, 0x94, 0x1C)
_COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)  # 纯黑色
_COLOR_GRAY = RGBColor(0x71, 0x80, 0x96)   # 页脚装饰用灰色


def _set_cell_shading(cell, color_hex: str):
    """设置单元格底色（已废弃，保留接口兼容性）"""
    pass


def _set_cell_text(cell, text: str, bold: bool = True, size: int = 12,
                   align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    """设置单元格文字"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = _FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)
    if color:
        run.font.color.rgb = color
    # 缩小行距
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def _set_cell_two_line(cell, line1: str, line2: str, bold1: bool = True,
                       size1: int = 12, size2: int = 10,
                       align=WD_ALIGN_PARAGRAPH.CENTER):
    """设置单元格双行文字（第一行size1，第二行size2灰色）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

    run1 = p.add_run(line1)
    run1.bold = bold1
    run1.font.size = Pt(size1)
    run1.font.name = _FONT
    run1._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)

    if line2:
        run2 = p.add_run(f"\n{line2}")
        run2.bold = True
        run2.font.size = Pt(size2)
        run2.font.name = _FONT
        # 表格内文字纯黑色
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)


def _set_cell_multiline(cell, text: str, bold: bool = True, size: int = 10,
                        align=WD_ALIGN_PARAGRAPH.LEFT):
    """设置单元格多行文字（保留 \\n 换行，用 add_break 插入 <w:br/>）。

    用于左上角分类标签（含 2-3 行的产品规格）。10号字。
    """
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.15

    lines = (text or "").split("\n")
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = _FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)
        if idx < len(lines) - 1:
            run.add_break()


def generate_production_order_docx(
    data: dict,
    page_size: str = "A4",
    orientation: str = "portrait",
) -> bytes:
    """生成生产订单 .docx 文件，返回 bytes。

    Args:
        data: get_production_order_print_data 返回的数据
        page_size: A4/A3/A5/B5
        orientation: portrait/landscape
    """
    doc = Document()

    # ── 页面设置（参照 reference.docx：A4横版+1.27cm边距）──
    section = doc.sections[0]
    w, h = PAGE_SIZES.get(page_size, PAGE_SIZES["A4"])

    # 自动判断：列数>10时强制横版
    column_defs = data.get("column_defs", [])
    num_data_cols = len(column_defs)
    if num_data_cols > 10 and orientation == "portrait":
        orientation = "landscape"

    if orientation == "landscape":
        w, h = h, w
    section.page_width = Cm(w)
    section.page_height = Cm(h)
    section.orientation = WD_ORIENT.LANDSCAPE if orientation == "landscape" else WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    header = data.get("header", {})
    sub_tables = data.get("sub_tables", [])
    column_defs = data.get("column_defs", [])
    rows = data.get("rows", [])
    col_totals = data.get("col_totals", {})
    weight_totals = data.get("weight_totals", {})
    weight_t_totals = data.get("weight_t_totals", {})

    # ── 标题 ──
    p_company = doc.add_paragraph()
    p_company.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_company = p_company.add_run(header.get("company_name", ""))
    run_company.bold = True
    run_company.font.size = Pt(16)
    run_company.font.name = _FONT
    run_company._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)

    # 副标题
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_sub = p_sub.add_run("生产订单 · 库存单")
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = _COLOR_GRAY
    run_sub.font.name = _FONT
    run_sub._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)

    # 订单信息 + 制单审核信息（合并为一行）
    creator = header.get("creator_name", "")
    reviewer = header.get("reviewer_name", "")
    print_date = header.get("print_date", "")

    info_parts = [f"订单编号：{header.get('order_no', '')}"]
    if header.get("batch_no"):
        info_parts.append(f"批次号：{header['batch_no']}")
    if header.get("remark"):
        info_parts.append(f"备注：{header['remark']}")
    # 追加制单/审核/日期到同一行
    info_parts.append(f"制单人：{creator}")
    info_parts.append(f"审核人：{reviewer}")
    info_parts.append(f"日期：{print_date}")

    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_info = p_info.add_run("　　".join(info_parts))
    run_info.font.size = Pt(9)
    run_info.font.name = _FONT
    run_info._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)

    if not column_defs or not rows:
        p_empty = doc.add_paragraph("该订单暂无明细数据")
        p_empty.alignment = WD_ALIGN_PARAGRAPH.CENTER
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── 按型号分类拆分的子表 ──────────────────────────────
    tables_to_render = []
    if sub_tables:
        for sub in sub_tables:
            sub_cdefs = sub.get("column_defs", [])
            sub_rows = sub.get("rows", [])
            sub_ct = sub.get("col_totals", {})
            sub_remarks = sub.get("category_remarks", [])
            sub_label = sub.get("category_label", "颜色")
            if sub_cdefs:
                tables_to_render.append((sub_cdefs, sub_rows, sub_ct, sub_remarks, sub_label))
    else:
        # 无分类数据时兜底：渲染全量单表
        tables_to_render.append((column_defs, rows, col_totals, [], "颜色"))

    # ── 动态列宽分配：确保表格适应页面宽度 ──
    usable_width_cm = w - section.left_margin.cm - section.right_margin.cm
    COLOR_COL_WIDTH = 3.5  # 颜色列固定宽度 cm
    TOTAL_COL_WIDTH = 2.0  # 合计列固定宽度 cm

    for tbl_idx, (tbl_cdefs, tbl_rows, tbl_ct, tbl_remarks, tbl_label) in enumerate(tables_to_render):
        if tbl_idx > 0:
            doc.add_paragraph()  # 子表间空行

        num_data_cols = len(tbl_cdefs)
        num_cols = 1 + num_data_cols + 1  # 颜色列 + 数据列 + 合计列

        # 计算数据列可用宽度，均分
        data_cols_width = usable_width_cm - COLOR_COL_WIDTH - TOTAL_COL_WIDTH
        if num_data_cols > 0:
            per_col_width = data_cols_width / num_data_cols
        else:
            per_col_width = 2.0

        total_doc_rows = 2 + len(tbl_rows) + 1
        table = doc.add_table(rows=total_doc_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # 设置表头行重复（跨页时每页显示表头）
        # 前两行是表头（一级group-header + 二级size-header）
        for row_idx in range(2):
            table.rows[row_idx]._element.get_or_add_trPr().append(
                OxmlElement('w:tblHeader')
            )

        # 设置列宽
        table.columns[0].width = Cm(COLOR_COL_WIDTH)
        for col_i in range(num_data_cols):
            table.columns[1 + col_i].width = Cm(per_col_width)
        table.columns[num_cols - 1].width = Cm(TOTAL_COL_WIDTH)

        # ── 一级表头（group-header）────────────────────────────
        cell_color_header = table.cell(0, 0)
        _set_cell_multiline(cell_color_header, tbl_label, bold=True, size=10,
                            align=WD_ALIGN_PARAGRAPH.LEFT)
        cell_color_header.merge(table.cell(1, 0))

        last_col_idx = num_cols - 1
        cell_total_header = table.cell(0, last_col_idx)
        _set_cell_text(cell_total_header, "合计", bold=True, size=12)
        cell_total_header.merge(table.cell(1, last_col_idx))

        # 合并 group 列
        col_idx = 1
        current_group = None
        for col_def in tbl_cdefs:
            if col_def["group"] != current_group:
                current_group = col_def["group"]
                group_cols = [c for c in tbl_cdefs if c["group"] == current_group]
                span = len(group_cols)
                start = col_idx
                end = col_idx + span - 1
                cell = table.cell(0, start)
                if span > 1:
                    cell.merge(table.cell(0, end))
                _set_cell_text(cell, current_group, bold=True, size=12)
            col_idx += 1

        # ── 二级表头（size-header）────────────────────────────
        for i, col_def in enumerate(tbl_cdefs):
            cell = table.cell(1, 1 + i)
            size_req = col_def.get("size_req", "")
            if size_req:
                _set_cell_two_line(cell, col_def["size"], size_req)
            else:
                _set_cell_text(cell, col_def["size"], size=12)

        # ── 数据行 ────────────────────────────────────────────
        for row_i, row_data in enumerate(tbl_rows):
            actual_row = 2 + row_i

            color_cell = table.cell(actual_row, 0)
            color_req = row_data.get("production_color_requirement", "")
            color_main = row_data.get("color", "")
            if color_req:
                _set_cell_two_line(color_cell, color_main, color_req,
                                   bold1=True, align=WD_ALIGN_PARAGRAPH.LEFT)
            else:
                _set_cell_text(color_cell, color_main, bold=True,
                               align=WD_ALIGN_PARAGRAPH.LEFT)

            for col_i, col_def in enumerate(tbl_cdefs):
                val = row_data.get(col_def["key"], 0)
                cell = table.cell(actual_row, 1 + col_i)
                # 0 值显示为空，所有文字纯黑色
                display_text = "" if val == 0 else str(val)
                _set_cell_text(cell, display_text, size=12, bold=True)

            row_total = row_data.get("row_total", 0)
            total_cell = table.cell(actual_row, last_col_idx)
            _set_cell_text(total_cell, str(row_total), bold=True, size=12)

        # ── 合计行 ────────────────────────────────────────────
        footer_row = 2 + len(tbl_rows)
        _set_cell_text(table.cell(footer_row, 0), "合计", bold=True, size=12,
                       align=WD_ALIGN_PARAGRAPH.LEFT)

        for col_i, col_def in enumerate(tbl_cdefs):
            val = tbl_ct.get(col_def["key"], 0)
            cell = table.cell(footer_row, 1 + col_i)
            # 0 值显示为空
            display_text = "" if val == 0 else str(val)
            _set_cell_text(cell, display_text, bold=True, size=12)

        grand_total = tbl_ct.get("grand_total", 0)
        gt_cell = table.cell(footer_row, last_col_idx)
        _set_cell_text(gt_cell, str(grand_total), bold=True, size=12,
                       color=_COLOR_ACCENT)

    # ── 公斤数统计表 ──────────────────────────────────────
    if weight_totals:
        doc.add_paragraph()  # 空行
        p_wt_title = doc.add_paragraph()
        run_wt = p_wt_title.add_run("公斤数统计（kg）")
        run_wt.bold = True
        run_wt.font.size = Pt(10)
        run_wt.font.name = _FONT
        run_wt._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)

        # 公斤数表基于全量 column_defs，动态列宽
        wt_num_cols = 1 + len(column_defs) + 1
        wt_last_col = wt_num_cols - 1

        # 列宽分配
        wt_data_cols_width = usable_width_cm - COLOR_COL_WIDTH - TOTAL_COL_WIDTH
        if len(column_defs) > 0:
            wt_per_col_width = wt_data_cols_width / len(column_defs)
        else:
            wt_per_col_width = 2.0

        # 公斤数表：theader 2 行 + 数据行(纯色 + T色)
        has_t = weight_t_totals.get("grand_total", 0) > 0
        wt_data_rows = 1 + (1 if has_t else 0)
        wt_table = doc.add_table(rows=2 + wt_data_rows, cols=wt_num_cols)
        wt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        wt_table.style = "Table Grid"

        # 设置表头行重复（跨页时每页显示表头）
        # 前两行是表头（一级等级header + 二级size-header）
        for row_idx in range(2):
            wt_table.rows[row_idx]._element.get_or_add_trPr().append(
                OxmlElement('w:tblHeader')
            )

        # 设置列宽
        wt_table.columns[0].width = Cm(COLOR_COL_WIDTH)
        for col_i in range(len(column_defs)):
            wt_table.columns[1 + col_i].width = Cm(wt_per_col_width)
        wt_table.columns[wt_last_col].width = Cm(TOTAL_COL_WIDTH)

        # 一级表头
        wt_cell0 = wt_table.cell(0, 0)
        _set_cell_text(wt_cell0, "等级", bold=True, size=12)
        wt_cell0.merge(wt_table.cell(1, 0))

        wt_cell_last = wt_table.cell(0, wt_last_col)
        _set_cell_text(wt_cell_last, "合计", bold=True, size=12)
        wt_cell_last.merge(wt_table.cell(1, wt_last_col))

        col_idx = 1
        current_group = None
        for col_def in column_defs:
            if col_def["group"] != current_group:
                current_group = col_def["group"]
                span = len([c for c in column_defs if c["group"] == current_group])
                start = col_idx
                end = col_idx + span - 1
                cell = wt_table.cell(0, start)
                if span > 1:
                    cell.merge(wt_table.cell(0, end))
                _set_cell_text(cell, current_group, bold=True, size=12)
            col_idx += 1

        # 二级表头
        for i, col_def in enumerate(column_defs):
            cell = wt_table.cell(1, 1 + i)
            size_req = col_def.get("size_req", "")
            if size_req:
                _set_cell_two_line(cell, col_def["size"], size_req)
            else:
                _set_cell_text(cell, col_def["size"], size=12)

        # 纯色行
        wt_row = 2
        _set_cell_text(wt_table.cell(wt_row, 0), "纯色", bold=True, size=12,
                       align=WD_ALIGN_PARAGRAPH.LEFT)
        for col_i, col_def in enumerate(column_defs):
            val = weight_totals.get(col_def["key"], 0)
            cell = wt_table.cell(wt_row, 1 + col_i)
            # 0 值显示为空，所有文字纯黑色
            display_text = "" if val == 0 else str(val)
            _set_cell_text(cell, display_text, size=12, bold=True)
        gt = weight_totals.get("grand_total", 0)
        _set_cell_text(wt_table.cell(wt_row, wt_last_col), str(gt),
                       bold=True, size=12, color=_COLOR_ACCENT)

        # T色行
        if has_t:
            wt_row += 1
            _set_cell_text(wt_table.cell(wt_row, 0), "T色", bold=True, size=12,
                           align=WD_ALIGN_PARAGRAPH.LEFT)
            for col_i, col_def in enumerate(column_defs):
                val = weight_t_totals.get(col_def["key"], 0)
                cell = wt_table.cell(wt_row, 1 + col_i)
                # 0 值显示为空，所有文字纯黑色
                display_text = "" if val == 0 else str(val)
                _set_cell_text(cell, display_text, size=12, bold=True)
            gt_t = weight_t_totals.get("grand_total", 0)
            _set_cell_text(wt_table.cell(wt_row, wt_last_col), str(gt_t),
                           bold=True, size=12, color=_COLOR_ACCENT)

    # ── 中间打孔要求图（订单含 model 关键词「中间打孔」时另起一页插入图片）──
    # 注意顺序：图片放在签字/页脚之前。否则当数据表正好填满页面时，签字+页脚
    # 会单独占下一页（只有 3 行内容显得空），再被 page_break_before 推到第三页放图片。
    has_punch = bool(data.get("has_middle_punch")) and _MIDDLE_PUNCH_IMAGE.exists()
    if has_punch:
        # 标题段：把分页符挂在段落属性上（page_break_before），避免 add_page_break() 插入额外空段落
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_title.paragraph_format.page_break_before = True
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(6)
        run_title = p_title.add_run("中间打孔要求")
        run_title.bold = True
        run_title.font.size = Pt(13)
        run_title.font.name = _FONT
        run_title._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)

        # 图片宽度按「同时满足页面可用宽度 + 可用高度」双约束，取较小值，保证单页放得下
        # 图片素材：1055 × 1491 px，高宽比 ≈ 1.413
        # 含签字/页脚预留：标题 ~1.2cm + 签字 ~1.0cm + 页脚 ~0.6cm + 段间距 ~0.4cm
        IMG_ASPECT_H_OVER_W = 1491 / 1055
        TITLE_RESERVE_CM = 1.2
        SIG_FOOTER_RESERVE_CM = 2.0  # 给图片下方的签字+页脚留位置，避免它们独占一页
        BOTTOM_BUFFER_CM = 0.3
        usable_w = max(w - 2.0, 6.0)  # 页宽 - 左右页边距 1.0+1.0
        usable_h = max(
            h - 1.2 - 1.0 - TITLE_RESERVE_CM - SIG_FOOTER_RESERVE_CM - BOTTOM_BUFFER_CM,
            6.0,
        )
        # 高度约束反推宽度：若按 usable_w 缩放后高度超过 usable_h，则改用高度限制
        width_limited_by_height = usable_h / IMG_ASPECT_H_OVER_W
        img_width_cm = min(usable_w, width_limited_by_height)

        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(0)
        p_img.add_run().add_picture(str(_MIDDLE_PUNCH_IMAGE), width=Cm(img_width_cm))

    # ── 页脚（仅页码：第X页/共X页，纯黑四号加粗居中） ──
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(0)
    footer_para.paragraph_format.space_after = Pt(0)

    # 页码：第X页/共X页（四号=14pt，纯黑色，加粗）
    run_page_text1 = footer_para.add_run("第")
    run_page_text1.font.size = Pt(14)
    run_page_text1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run_page_text1.font.name = _FONT
    run_page_text1.bold = True
    run_page_text1._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)

    # 当前页码（域代码）
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run_page_num = footer_para.add_run()
    run_page_num._element.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run_page_num._element.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_page_num._element.append(fldChar2)
    run_page_num.font.size = Pt(14)
    run_page_num.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run_page_num.font.name = _FONT
    run_page_num.bold = True

    run_page_text2 = footer_para.add_run("页/共")
    run_page_text2.font.size = Pt(14)
    run_page_text2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run_page_text2.font.name = _FONT
    run_page_text2.bold = True
    run_page_text2._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)

    # 总页数（域代码）
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run_total_pages = footer_para.add_run()
    run_total_pages._element.append(fldChar3)

    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = "NUMPAGES"
    run_total_pages._element.append(instrText2)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run_total_pages._element.append(fldChar4)
    run_total_pages.font.size = Pt(14)
    run_total_pages.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run_total_pages.font.name = _FONT
    run_total_pages.bold = True

    run_page_text3 = footer_para.add_run("页")
    run_page_text3.font.size = Pt(14)
    run_page_text3.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run_page_text3.font.name = _FONT
    run_page_text3.bold = True
    run_page_text3._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
