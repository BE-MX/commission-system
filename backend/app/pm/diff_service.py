"""PM Hub AI 差异概要管线：本地先算精确 diff，AI 只把 diff 讲成人话。

不把两份全文塞给模型（价格体系上的幻觉差异比没有概要更危险）：
  抽取本版 & 上一版文本 → 本地 diff（文本: difflib；xlsx: openpyxl 单元格级）
  → diff 喂给平台 ai.service（pm_diff preset）生成概要 → 存版本记录

降级策略：v1 无上一版 → not_applicable；扫描版 PDF / 图片压缩包 → not_applicable；
AI 失败 → failed 可手动重试，不影响版本保存。pending 看门狗在启动时回收
（服务重启丢在途后台任务，展会 STALE_PENDING 同款做法）。
"""

import difflib
import logging
from pathlib import Path
from typing import Optional

from sqlalchemy.orm import Session

from app.core.database import SessionLocal
from app.pm.material_service import get_version, previous_version
from app.pm.models import PmMaterialVersion, bj_now
from app.pm.service import to_abs

logger = logging.getLogger("commission")

DIFF_PRESET = "pm_diff"
MAX_DIFF_CHARS = 12000  # 喂给模型的 diff 安全长度
MAX_XLSX_CHANGES = 300  # xlsx 单元格级变更条数上限
STALE_DIFF_SECS = 600  # pending 看门狗阈值（必须大于 AI 调用最坏墙钟）

TEXT_EXTS = {".md", ".markdown", ".txt", ".csv", ".log", ".json"}
DOCX_EXTS = {".docx"}
PDF_EXTS = {".pdf"}
XLSX_EXTS = {".xlsx", ".xlsm"}

DIFF_SYSTEM_PROMPT = """你是项目资料版本差异的整理助手。输入是两版文件经本地精确比对产生的 diff 结果（unified diff 或单元格变更清单），你的唯一任务是把它转述成人话。

【输出格式】
按「新增 / 修改 / 删除」分类列出 3~8 条要点，每条一行、以「· 」开头；最后空一行写一句话总评（以「总评：」开头）。没有变更的分类直接省略。

【规则】
- 仅陈述差异本身，不评价内容好坏，不推测意图
- 数字、价格、比例、日期必须照抄 diff 中的原始值，禁止约算或编造
- diff 中不存在的变更禁止杜撰
- 要点要具体到条目（如「L1+ 客户折扣从 8% 调整为 8-12% 区间」），不要「部分内容有调整」这类空话
"""


# ── 文本抽取 ─────────────────────────────────────────────────────────

def extract_text(abs_path: Path, ext: str) -> Optional[str]:
    """抽取纯文本。返回 None = 不支持/抽取为空（调用方落 not_applicable）。"""
    try:
        if ext in TEXT_EXTS:
            return abs_path.read_text(encoding="utf-8", errors="replace")
        if ext in DOCX_EXTS:
            from docx import Document
            doc = Document(str(abs_path))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        if ext in PDF_EXTS:
            from pypdf import PdfReader
            reader = PdfReader(str(abs_path))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
            # 扫描版 PDF 抽出为空——不拿空文本硬比
            return text if text.strip() else None
    except Exception as exc:
        logger.warning("[PM] extract_text failed %s: %s", abs_path, exc)
        print(f"[PM] extract_text failed {abs_path}: {exc}", flush=True)
        return None
    return None


# ── 本地 diff ────────────────────────────────────────────────────────

def diff_text(old: str, new: str) -> tuple[str, bool]:
    """unified diff。返回 (diff 文本, 是否截断)。"""
    old_lines = old.splitlines()
    new_lines = new.splitlines()
    diff_lines = list(difflib.unified_diff(old_lines, new_lines, lineterm="", n=1))
    truncated = False
    text = "\n".join(diff_lines)
    if len(text) > MAX_DIFF_CHARS:
        text = text[:MAX_DIFF_CHARS]
        truncated = True
    return text, truncated


def _sheet_cells(ws) -> dict:
    cells = {}
    for row in ws.iter_rows():
        for cell in row:
            if cell.value is not None:
                cells[(cell.row, cell.column)] = str(cell.value)
    return cells


def diff_xlsx(old_path: Path, new_path: Path) -> tuple[str, bool]:
    """openpyxl 行/单元格级结构对比（data_only=True 取公式缓存值，全部 sheet）。"""
    import openpyxl
    old_wb = openpyxl.load_workbook(str(old_path), data_only=True, read_only=True)
    new_wb = openpyxl.load_workbook(str(new_path), data_only=True, read_only=True)
    changes: list[str] = []
    truncated = False
    sheet_names = list(dict.fromkeys(old_wb.sheetnames + new_wb.sheetnames))
    for name in sheet_names:
        if name not in old_wb.sheetnames:
            changes.append(f"[工作表] 新增「{name}」")
            continue
        if name not in new_wb.sheetnames:
            changes.append(f"[工作表] 删除「{name}」")
            continue
        old_cells = _sheet_cells(old_wb[name])
        new_cells = _sheet_cells(new_wb[name])
        for key in sorted(set(old_cells) | set(new_cells)):
            r, c = key
            old_v, new_v = old_cells.get(key), new_cells.get(key)
            if old_v == new_v:
                continue
            cell_ref = f"{name}!R{r}C{c}"
            if old_v is None:
                changes.append(f"[新增] {cell_ref} = {new_v}")
            elif new_v is None:
                changes.append(f"[删除] {cell_ref} 原值 {old_v}")
            else:
                changes.append(f"[修改] {cell_ref}: {old_v} → {new_v}")
            if len(changes) >= MAX_XLSX_CHANGES:
                truncated = True
                break
        if truncated:
            break
    old_wb.close()
    new_wb.close()
    return "\n".join(changes), truncated


# ── 主管线 ───────────────────────────────────────────────────────────

def generate_diff_summary(db: Session, version_id: int) -> None:
    """计算并写回某个版本的 AI 差异概要。任何失败只落状态，不抛出。"""
    version = get_version(db, version_id)
    if not version or version.deleted_at is not None:
        return
    try:
        _generate(db, version)
    except Exception as exc:  # 单条版本 AI 失败不影响上传主流程（宪法 6：log+print，不无声吞）
        db.rollback()
        logger.warning("[PM] diff pipeline failed version=%s: %s", version_id, exc)
        print(f"[PM] diff pipeline failed version={version_id}: {exc}", flush=True)
        try:
            version = get_version(db, version_id)
            if version:
                version.diff_status = "failed"
                version.diff_error = str(exc)[:500]
                version.diff_updated_at = bj_now()
                db.commit()
        except Exception as exc2:
            db.rollback()
            logger.warning("[PM] diff fail-status writeback failed version=%s: %s", version_id, exc2)
            print(f"[PM] diff fail-status writeback failed version={version_id}: {exc2}", flush=True)


def _generate(db: Session, version: PmMaterialVersion) -> None:
    ext = Path(version.original_name).suffix.lower()
    prev = previous_version(db, version)
    if not prev:
        version.diff_status = "not_applicable"
        version.diff_summary = None
        version.diff_updated_at = bj_now()
        db.commit()
        return

    new_path = to_abs(version.file_path)
    old_path = to_abs(prev.file_path)

    if ext in XLSX_EXTS:
        diff_text_str, truncated = diff_xlsx(old_path, new_path)
        if not diff_text_str.strip():
            version.diff_status = "done"
            version.diff_summary = "两版内容完全一致（无单元格级变更）。"
            version.diff_updated_at = bj_now()
            db.commit()
            return
    elif ext in TEXT_EXTS | DOCX_EXTS | PDF_EXTS:
        old_text = extract_text(old_path, ext)
        new_text = extract_text(new_path, ext)
        if old_text is None or new_text is None:
            version.diff_status = "not_applicable"
            version.diff_summary = None
            version.diff_error = "扫描件或无法抽取文本"
            version.diff_updated_at = bj_now()
            db.commit()
            return
        diff_text_str, truncated = diff_text(old_text, new_text)
        if not diff_text_str.strip():
            version.diff_status = "done"
            version.diff_summary = "两版文本内容完全一致。"
            version.diff_updated_at = bj_now()
            db.commit()
            return
    else:
        version.diff_status = "not_applicable"
        version.diff_summary = None
        version.diff_error = "该类型不支持对比"
        version.diff_updated_at = bj_now()
        db.commit()
        return

    from app.ai.service import chat  # 延迟 import，避免模块装载期循环
    from app.pm.material_service import get_material
    material = get_material(db, version.material_id, include_deleted=True)
    material_name = material.name if material else str(version.material_id)
    change_count = diff_text_str.count("\n") + 1
    header = f"变更过多，仅摘要前 {change_count} 处\n" if truncated else ""
    user_msg = (
        f"资料《{material_name}》v{prev.version_no} → v{version.version_no} 的本地精确 diff 如下，"
        f"请按约定格式转述：\n\n{header}{diff_text_str}"
    )
    result = chat(db, DIFF_PRESET, [{"role": "user", "content": user_msg}], caller_module="pm")
    summary = (result.get("content") or "").strip()
    if truncated:
        summary = f"⚠ 变更过多，仅摘要前 {change_count} 处\n" + summary
    version.diff_status = "done"
    version.diff_summary = summary
    version.diff_error = None
    version.diff_updated_at = bj_now()
    db.commit()


def run_diff_in_background(version_id: int) -> None:
    """BackgroundTask 入口：自建 SessionLocal + finally 关闭。

    chat() 需要活 session，而请求级 get_db 在后台任务执行前已关闭——
    归平台红线 4 的线程池许可场景（aftersales 同款先例）。
    """
    db = SessionLocal()
    try:
        generate_diff_summary(db, version_id)
    finally:
        db.close()


def recover_stale_diffs() -> None:
    """启动时回收超时 pending（服务重启丢在途后台任务）。warn 不阻塞启动。"""
    try:
        from datetime import timedelta
        db = SessionLocal()
        try:
            threshold = bj_now() - timedelta(seconds=STALE_DIFF_SECS)
            stale = (
                db.query(PmMaterialVersion)
                .filter(
                    PmMaterialVersion.diff_status == "pending",
                    PmMaterialVersion.created_at < threshold,
                )
                .all()
            )
            for v in stale:
                v.diff_status = "failed"
                v.diff_error = f"服务重启回收：pending 超过 {STALE_DIFF_SECS}s"
                v.diff_updated_at = bj_now()
            if stale:
                db.commit()
                logger.info("[PM] recovered %d stale diff tasks", len(stale))
                print(f"[PM] recovered {len(stale)} stale diff tasks", flush=True)
        finally:
            db.close()
    except Exception as exc:
        logger.warning("[PM] recover_stale_diffs skipped: %s", exc)
        print(f"[PM] recover_stale_diffs skipped: {exc}", flush=True)
