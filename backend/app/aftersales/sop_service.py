"""售后 SOP 版本、结构化解析和启用。"""

from datetime import date, datetime
from hashlib import sha256
from pathlib import Path

from docx import Document
from sqlalchemy.orm import Session

from app.aftersales.file_service import resolve_private_path, store_bytes, validate_upload
from app.aftersales.models import AfterSalesEvent, AfterSalesSopVersion
from app.auth.models import ArkUser

try:
    from pypdf import PdfReader

    _PDF_READER_AVAILABLE = True
except ImportError:  # pragma: no cover - deployment dependency determines this branch
    PdfReader = None
    _PDF_READER_AVAILABLE = False


class SopParseError(ValueError):
    pass


ISSUE_TITLES = {
    "断发": ("断发",),
    "褪色": ("褪色", "变色"),
    "色差": ("颜色色差", "色差"),
    "头发太直": ("头发太直", "不够自然"),
    "脱发": ("脱发", "掉发"),
    "分叉": ("分叉",),
    "发干打结": ("发干", "打结"),
    "头发油": ("头发油", "油腻"),
    "贴发胶": ("贴发胶",),
    "克重不够": ("克重不够",),
    "产品做工": ("抹胶", "包装细节", "产品做工"),
}


def _issue_mapping(section_titles: list[str]) -> dict[str, str | list[str]]:
    mapping = {}
    for issue, keywords in ISSUE_TITLES.items():
        matched = [title for title in section_titles if any(keyword in title for keyword in keywords)]
        if len(matched) == 1:
            mapping[issue] = matched[0]
        elif matched:
            mapping[issue] = matched
    return mapping


def _parse_docx(path: Path) -> dict:
    try:
        document = Document(path)
    except Exception as exc:
        raise SopParseError("DOCX 文件无法读取或已经损坏") from exc

    sections: list[dict] = []
    current = None
    loose_paragraphs: list[str] = []
    clause_count = 0
    for paragraph in document.paragraphs:
        content = paragraph.text.strip()
        if not content:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading") or style_name.startswith("标题"):
            level_digits = "".join(character for character in style_name if character.isdigit())
            current = {
                "title": content,
                "level": int(level_digits or 1),
                "paragraphs": [],
            }
            sections.append(current)
        elif current is None:
            loose_paragraphs.append(content)
            clause_count += 1
        else:
            current["paragraphs"].append(content)
            clause_count += 1

    tables = []
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables.append({"rows": rows})
        clause_count += max(0, len(rows) - 1)

    titles = [section["title"] for section in sections]
    return {
        "loose_paragraphs": loose_paragraphs,
        "sections": sections,
        "tables": tables,
        "issue_mapping": _issue_mapping(titles),
        "clause_count": clause_count,
    }


def _parse_pdf(path: Path) -> dict:
    if not _PDF_READER_AVAILABLE:
        raise SopParseError("PDF 解析需要安装 pypdf，请联系管理员补齐依赖后重试")
    try:
        paragraphs = [
            line.strip()
            for page in PdfReader(str(path)).pages
            for line in (page.extract_text() or "").splitlines()
            if line.strip()
        ]
    except Exception as exc:
        raise SopParseError("PDF 文件无法读取或没有可提取文本") from exc
    if not paragraphs:
        raise SopParseError("PDF 没有可提取文本，请上传可搜索 PDF 或 DOCX")
    return {
        "loose_paragraphs": paragraphs,
        "sections": [],
        "tables": [],
        "issue_mapping": {},
        "clause_count": len(paragraphs),
    }


def parse_sop_file(path: str | Path) -> dict:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return _parse_docx(file_path)
    if suffix == ".pdf":
        return _parse_pdf(file_path)
    raise SopParseError("SOP 仅支持 DOCX 或 PDF")


def create_sop_version(
    db: Session,
    *,
    storage_root: str | Path,
    original_filename: str,
    content: bytes,
    version_no: str,
    change_summary: str,
    effective_date: date,
    uploaded_by_user_id: int,
) -> AfterSalesSopVersion:
    if db.query(AfterSalesSopVersion.id).filter(AfterSalesSopVersion.version_no == version_no).first():
        raise ValueError("SOP 版本号已存在")
    mime_type = (
        "application/pdf"
        if original_filename.lower().endswith(".pdf")
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    validate_upload(original_filename, mime_type, len(content), purpose="sop")
    storage_path = store_bytes(storage_root, "sop", original_filename, content)
    parsed = parse_sop_file(resolve_private_path(storage_root, storage_path))
    version = AfterSalesSopVersion(
        version_no=version_no,
        original_filename=original_filename,
        storage_path=storage_path,
        file_hash=sha256(content).hexdigest(),
        change_summary=change_summary,
        effective_date=effective_date,
        parse_status="parsed",
        structured_content_json=parsed,
        issue_mapping_json=parsed["issue_mapping"],
        clause_count=parsed["clause_count"],
        uploaded_by_user_id=uploaded_by_user_id,
    )
    db.add(version)
    db.flush()
    return version


def activate_sop_version(
    db: Session,
    version_id: int,
    actor_user_id: int,
) -> AfterSalesSopVersion:
    version = db.query(AfterSalesSopVersion).filter(AfterSalesSopVersion.id == version_id).first()
    if version is None:
        raise ValueError("SOP 版本不存在")
    if version.parse_status != "parsed":
        raise ValueError("SOP 尚未解析完成，不能启用")
    db.query(AfterSalesSopVersion).filter(AfterSalesSopVersion.is_active.is_(True)).update(
        {AfterSalesSopVersion.is_active: False}, synchronize_session=False
    )
    version.is_active = True
    version.activated_at = datetime.utcnow()
    actor = db.get(ArkUser, actor_user_id)
    db.add(
        AfterSalesEvent(
            case_id=None,
            event_type="sop_activated",
            actor_user_id=actor_user_id,
            actor_name_snapshot=actor.real_name if actor else None,
            detail_json={
                "sop_version_id": version.id,
                "version_no": version.version_no,
                "change_summary": version.change_summary,
            },
        )
    )
    db.flush()
    return version
