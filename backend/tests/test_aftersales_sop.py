from datetime import date
from io import BytesIO

import pytest
from docx import Document

from app.aftersales.file_service import (
    FileValidationError,
    resolve_private_path,
    store_bytes,
    validate_upload,
)
from app.aftersales.models import AfterSalesEvent, AfterSalesSopVersion
from app.aftersales.sop_service import (
    SopParseError,
    activate_sop_version,
    create_sop_version,
    parse_sop_file,
)
from app.auth.models import ArkUser


def _docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("统一售后处理流程", level=1)
    doc.add_heading("初步判责", level=2)
    doc.add_paragraph("A 类：明确生产或包装问题。")
    doc.add_paragraph("D 类：责任暂不明确，先收集证据。")
    doc.add_heading("褪色 / 变色问题", level=1)
    doc.add_paragraph("处理原则：区分 fading 与 color changing。")
    doc.add_paragraph("优化话术：Thank you for letting us know.")
    doc.add_heading("贴发 / 天才发帘抹胶厚薄不匀", level=1)
    doc.add_paragraph("轻微差异属于手工生产正常范围。")
    doc.add_heading("包装细节出错", level=1)
    doc.add_paragraph("包装错误应主动承担并补发正确包装。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "售后阶段"
    table.cell(0, 1).text = "输出结果"
    table.cell(1, 0).text = "收集证据"
    table.cell(1, 1).text = "形成判断依据"
    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


def _user(db, username="sop-admin"):
    user = ArkUser(
        username=username,
        password_hash="test-hash",
        real_name="SOP Admin",
        dingtalk_id=f"ding-{username}",
    )
    db.add(user)
    db.flush()
    return user


def test_docx_parser_keeps_headings_paragraphs_tables_and_issue_mapping(tmp_path):
    path = tmp_path / "售后SOP.docx"
    path.write_bytes(_docx_bytes())

    parsed = parse_sop_file(path)

    assert parsed["sections"][0]["title"] == "统一售后处理流程"
    assert any(section["title"] == "褪色 / 变色问题" for section in parsed["sections"])
    assert parsed["tables"][0]["rows"][1] == ["收集证据", "形成判断依据"]
    assert parsed["issue_mapping"]["褪色"] == "褪色 / 变色问题"
    assert parsed["issue_mapping"]["产品做工"] == ["贴发 / 天才发帘抹胶厚薄不匀", "包装细节出错"]
    assert parsed["clause_count"] >= 4


def test_pdf_without_parser_has_actionable_error(tmp_path, monkeypatch):
    path = tmp_path / "售后SOP.pdf"
    path.write_bytes(b"%PDF-1.4 fake")
    monkeypatch.setattr("app.aftersales.sop_service._PDF_READER_AVAILABLE", False)

    with pytest.raises(SopParseError, match="pypdf"):
        parse_sop_file(path)


def test_upload_validation_enforces_extension_and_size():
    validate_upload("evidence.jpg", "image/jpeg", 20 * 1024 * 1024, purpose="evidence")
    with pytest.raises(FileValidationError, match="格式"):
        validate_upload("evidence.exe", "application/octet-stream", 10, purpose="evidence")
    with pytest.raises(FileValidationError, match="20MB"):
        validate_upload("large.jpg", "image/jpeg", 20 * 1024 * 1024 + 1, purpose="evidence")


def test_storage_path_cannot_escape_root(tmp_path):
    with pytest.raises(FileValidationError, match="非法"):
        resolve_private_path(tmp_path, "../secret.txt")


def test_store_bytes_uses_random_private_path(tmp_path):
    first = store_bytes(tmp_path, "evidence", "客户照片.jpg", b"one")
    second = store_bytes(tmp_path, "evidence", "客户照片.jpg", b"two")

    assert first != second
    assert resolve_private_path(tmp_path, first).read_bytes() == b"one"
    assert "客户照片" not in first


def test_only_one_sop_version_can_be_active(db, tmp_path):
    admin = _user(db)
    first = create_sop_version(
        db,
        storage_root=tmp_path,
        original_filename="first.docx",
        content=_docx_bytes(),
        version_no="2026.07.10-v1",
        change_summary="首版",
        effective_date=date(2026, 7, 10),
        uploaded_by_user_id=admin.id,
    )
    second = create_sop_version(
        db,
        storage_root=tmp_path,
        original_filename="second.docx",
        content=_docx_bytes(),
        version_no="2026.07.10-v2",
        change_summary="修订版",
        effective_date=date(2026, 7, 11),
        uploaded_by_user_id=admin.id,
    )
    db.commit()

    activate_sop_version(db, first.id, admin.id)
    db.commit()
    activate_sop_version(db, second.id, admin.id)
    db.commit()

    versions = db.query(AfterSalesSopVersion).order_by(AfterSalesSopVersion.id).all()
    assert [item.is_active for item in versions] == [False, True]
    assert versions[1].activated_at is not None
    audits = db.query(AfterSalesEvent).filter_by(event_type="sop_activated").all()
    assert len(audits) == 2
    assert audits[-1].case_id is None
    assert audits[-1].detail_json["sop_version_id"] == second.id


def test_duplicate_sop_version_number_is_rejected(db, tmp_path):
    admin = _user(db, "sop-admin-duplicate")
    common = dict(
        db=db,
        storage_root=tmp_path,
        original_filename="sop.docx",
        content=_docx_bytes(),
        version_no="2026.07.10-v1",
        change_summary="首版",
        effective_date=date(2026, 7, 10),
        uploaded_by_user_id=admin.id,
    )
    create_sop_version(**common)
    db.commit()

    with pytest.raises(ValueError, match="版本号"):
        create_sop_version(**common)
