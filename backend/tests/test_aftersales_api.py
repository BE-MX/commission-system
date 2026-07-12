from contextlib import contextmanager
from datetime import date, datetime, timedelta
from decimal import Decimal
from io import BytesIO

from fastapi import FastAPI
from fastapi.testclient import TestClient
from docx import Document
from sqlalchemy import text

from app.aftersales.models import AfterSalesCase, AfterSalesNotificationLog
from app.auth.models import ArkUser, ArkUserExternalBinding
from app.auth.utils import create_access_token
from app.core.database import get_db
from app.models.employee import SupervisorRelationHistory


def _user(db, username="api-sales"):
    user = ArkUser(
        username=username,
        password_hash="test-hash",
        real_name="API Sales",
        dingtalk_id=f"ding-{username}",
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return user


@contextmanager
def _client(db, user, permissions):
    from app.aftersales.router import router

    app = FastAPI()
    app.include_router(router, prefix="/api/aftersales")

    def override_db():
        yield db

    app.dependency_overrides[get_db] = override_db
    token = create_access_token(
        {
            "sub": str(user.id),
            "username": user.username,
            "real_name": user.real_name,
            "roles": [],
            "permissions": permissions,
        }
    )
    with TestClient(app, headers={"Authorization": f"Bearer {token}"}) as client:
        yield client


def _payload():
    return {
        "customer_id": "CUST001",
        "customer_name_snapshot": "客户A",
        "customer_grade": "A",
        "order_id": "ORD001",
        "order_no_snapshot": "NO001",
        "purchase_date": "2026-07-01",
        "feedback_date": "2026-07-10",
        "product_id": 1,
        "product_name_snapshot": "Invisible Weft",
        "is_custom_product": False,
        "batch_no": "BATCH-1",
        "color_value": "#2B",
        "length_value": "20 inch",
        "weight_value": "100",
        "weight_unit": "g",
        "quantity": "2",
        "primary_issue_type": "褪色",
        "problem_description": "客户使用三周后出现明显褪色，已经影响终端客户继续销售。",
        "occurred_stage": "使用几天",
        "care_storage_note": "客户使用无硫酸盐洗发水，低温护理，未游泳或暴晒。",
        "affects_end_customer": "yes",
        "affected_goods_value": "1150",
        "affected_goods_currency": "USD",
        "sales_evidence_confirmed": True,
    }


def test_create_and_get_case_use_response_envelope(db, monkeypatch, tmp_path):
    monkeypatch.setattr("app.aftersales.router._storage_root", lambda: tmp_path)
    user = _user(db)
    with _client(db, user, ["aftersales:read", "aftersales:write"]) as client:
        created = client.post("/api/aftersales/cases", json=_payload())
        assert created.status_code == 200
        body = created.json()
        assert body["code"] == 200
        assert body["data"]["current_status"] == "draft"

        detail = client.get(f"/api/aftersales/cases/{body['data']['id']}")
        assert detail.status_code == 200
        assert detail.json()["data"]["customer_name_snapshot"] == "客户A"


def test_create_requires_write_permission(db):
    user = _user(db, "api-readonly")
    with _client(db, user, ["aftersales:read"]) as client:
        response = client.post("/api/aftersales/cases", json=_payload())

    assert response.status_code == 403


def test_case_list_defaults_to_own_scope(db):
    own = _user(db, "api-own")
    other = _user(db, "api-other")
    with _client(db, own, ["aftersales:read", "aftersales:write"]) as client:
        client.post("/api/aftersales/cases", json=_payload())
    with _client(db, other, ["aftersales:read", "aftersales:write"]) as client:
        client.post("/api/aftersales/cases", json={**_payload(), "order_id": "ORD002", "order_no_snapshot": "NO002"})

    with _client(db, own, ["aftersales:read"]) as client:
        response = client.get("/api/aftersales/cases")

    assert response.status_code == 200
    assert response.json()["data"]["total"] == 1
    assert response.json()["data"]["items"][0]["creator_user_id"] == own.id


def test_read_all_permission_can_see_all_cases(db):
    first = _user(db, "api-all-first")
    second = _user(db, "api-all-second")
    with _client(db, first, ["aftersales:read", "aftersales:write"]) as client:
        client.post("/api/aftersales/cases", json=_payload())
    with _client(db, second, ["aftersales:read", "aftersales:write"]) as client:
        client.post("/api/aftersales/cases", json={**_payload(), "order_id": "ORD003", "order_no_snapshot": "NO003"})

    with _client(db, first, ["aftersales:read", "aftersales:read_all"]) as client:
        response = client.get("/api/aftersales/cases?scope=all")

    assert response.json()["data"]["total"] == 2


def test_evidence_upload_is_private_and_creator_only(db, monkeypatch, tmp_path):
    monkeypatch.setattr("app.aftersales.router._storage_root", lambda: tmp_path)
    creator = _user(db, "api-upload")
    stranger = _user(db, "api-stranger")
    with _client(db, creator, ["aftersales:read", "aftersales:write"]) as client:
        case_id = client.post("/api/aftersales/cases", json=_payload()).json()["data"]["id"]
        uploaded = client.post(
            f"/api/aftersales/cases/{case_id}/evidence",
            data={"evidence_type": "overview_image"},
            files={"file": ("overview.jpg", b"image-bytes", "image/jpeg")},
        )
        assert uploaded.status_code == 200
        evidence = uploaded.json()["data"]
        assert "storage_path" not in evidence
        download = client.get(evidence["download_url"])
        assert download.status_code == 200
        assert download.content == b"image-bytes"

    with _client(db, stranger, ["aftersales:read", "aftersales:write"]) as client:
        forbidden = client.get(evidence["download_url"])
    assert forbidden.status_code == 403


def test_options_expose_complete_issue_and_action_dictionaries(db):
    user = _user(db, "api-options")
    with _client(db, user, ["aftersales:read"]) as client:
        response = client.get("/api/aftersales/options")

    data = response.json()["data"]
    assert len(data["issue_types"]) == 11
    assert any(item["code"] == "replacement" and item["has_compensation"] for item in data["actions"])


def test_customer_order_and_product_selector_endpoints(db, seed_business_data):
    db.execute(
        text(
            "CREATE TABLE lsordertest.okki_products ("
            "product_id INTEGER PRIMARY KEY, name TEXT, model TEXT, color TEXT, size TEXT, unit TEXT, disable_flag INTEGER)"
        )
    )
    db.execute(text("INSERT INTO lsordertest.okki_products VALUES (1, 'Invisible Weft', 'IW', '#2B', '20 inch', '100g', 0)"))
    db.commit()
    user = _user(db, "api-selectors")

    with _client(db, user, ["aftersales:read"]) as client:
        customers = client.get("/api/aftersales/customers/search?keyword=客户A").json()["data"]["items"]
        orders = client.get("/api/aftersales/orders/search?customer_id=CUST001").json()["data"]["items"]
        products = client.get("/api/aftersales/products/search?keyword=Weft").json()["data"]["items"]

    assert customers[0]["customer_id"] == "CUST001"
    assert {item["customer_id"] for item in orders} == {"CUST001"}
    assert products[0]["product_name"] == "Invisible Weft"


def _sop_docx():
    doc = Document()
    doc.add_heading("褪色 / 变色问题", level=1)
    doc.add_paragraph("处理原则：区分 fading 与 color changing。")
    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


def test_admin_can_upload_preview_and_activate_sop(db, monkeypatch, tmp_path):
    monkeypatch.setattr("app.aftersales.router._storage_root", lambda: tmp_path)
    admin = _user(db, "api-sop-admin")
    with _client(db, admin, ["aftersales:admin"]) as client:
        upload = client.post(
            "/api/aftersales/sop/versions",
            data={
                "version_no": "2026.07.10",
                "change_summary": "首版",
                "effective_date": "2026-07-10",
            },
            files={
                "file": (
                    "售后SOP.docx",
                    _sop_docx(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert upload.status_code == 200
        version = upload.json()["data"]
        assert version["parse_status"] == "parsed"
        assert version["issue_mapping"]["褪色"] == "褪色 / 变色问题"

        activated = client.post(f"/api/aftersales/sop/versions/{version['id']}/activate")
        assert activated.status_code == 200
        assert activated.json()["data"]["is_active"] is True


def test_http_decision_submit_and_supervisor_approval_flow(db, monkeypatch, tmp_path):
    monkeypatch.setattr("app.aftersales.router._storage_root", lambda: tmp_path)
    async def skip_delivery(_notification_ids):
        return None
    monkeypatch.setattr("app.aftersales.router._deliver_notification_task", skip_delivery)
    sales = _user(db, "api-flow-sales")
    supervisor = _user(db, "api-flow-supervisor")
    db.add_all(
        [
            ArkUserExternalBinding(
                ark_user_id=sales.id,
                provider="okki",
                external_account_id="SP-FLOW",
                binding_status="active",
                is_primary=True,
            ),
            ArkUserExternalBinding(
                ark_user_id=supervisor.id,
                provider="okki",
                external_account_id="SV-FLOW",
                binding_status="active",
                is_primary=True,
            ),
            SupervisorRelationHistory(
                salesperson_id="SP-FLOW",
                supervisor_id="SV-FLOW",
                effective_start=date(2026, 1, 1),
                is_current=True,
            ),
        ]
    )
    db.commit()

    with _client(db, sales, ["aftersales:read", "aftersales:write"]) as client:
        case_id = client.post("/api/aftersales/cases", json={**_payload(), "batch_no": None}).json()["data"]["id"]
        case = db.get(AfterSalesCase, case_id)
        case.current_status = "awaiting_sales_decision"
        db.commit()
        for evidence_type in ("overview_image", "closeup_image"):
            response = client.post(
                f"/api/aftersales/cases/{case_id}/evidence",
                data={"evidence_type": evidence_type},
                files={"file": (f"{evidence_type}.jpg", b"image", "image/jpeg")},
            )
            assert response.status_code == 200

        decision = client.post(
            f"/api/aftersales/cases/{case_id}/decision",
            json={
                "responsibility_class": "D",
                "responsibility_reason": "责任暂不明确",
                "actions": [{
                    "code": "return_inspection",
                    "freight_payer": "customer",
                    "return_address": "LeShine Quality Center",
                    "expected_completion_date": "2026-07-20",
                }],
                "customer_reply_draft": "Thank you for your feedback. Please return the affected hair for inspection.",
                "requires_return": True,
            },
        )
        assert decision.status_code == 200
        version = decision.json()["data"]["version"]
        submitted = client.post(
            f"/api/aftersales/cases/{case_id}/submit",
            json={"version": version, "idempotency_key": "api-submit-flow"},
        )
        assert submitted.status_code == 200
        assert submitted.json()["data"]["current_status"] == "awaiting_supervisor"
        version = submitted.json()["data"]["version"]

    with _client(db, supervisor, ["aftersales:read", "aftersales:write"]) as client:
        reviewed = client.post(
            f"/api/aftersales/cases/{case_id}/review",
            json={
                "decision": "approve",
                "comment": "同意",
                "version": version,
                "idempotency_key": "api-review-flow",
            },
        )

    assert reviewed.status_code == 200
    assert reviewed.json()["data"]["current_status"] == "approved"
    assert db.query(AfterSalesNotificationLog).filter_by(case_id=case_id).count() == 2


def test_deleting_evidence_recalculates_score_and_missing_items(db, monkeypatch, tmp_path):
    monkeypatch.setattr("app.aftersales.router._storage_root", lambda: tmp_path)
    user = _user(db, "api-evidence-delete")
    with _client(db, user, ["aftersales:read", "aftersales:write"]) as client:
        case_id = client.post(
            "/api/aftersales/cases", json={**_payload(), "batch_no": None}
        ).json()["data"]["id"]
        evidence_ids = []
        for evidence_type in ("overview_image", "closeup_image"):
            response = client.post(
                f"/api/aftersales/cases/{case_id}/evidence",
                data={"evidence_type": evidence_type},
                files={"file": (f"{evidence_type}.jpg", b"image", "image/jpeg")},
            )
            evidence_ids.append(response.json()["data"]["id"])
        before = client.get(f"/api/aftersales/cases/{case_id}").json()["data"]
        assert before["evidence_is_sufficient"] is True

        deleted = client.delete(
            f"/api/aftersales/cases/{case_id}/evidence/{evidence_ids[1]}"
        )
        after = client.get(f"/api/aftersales/cases/{case_id}").json()["data"]

    assert deleted.status_code == 200
    assert after["evidence_is_sufficient"] is False
    assert "问题近景图" in after["evidence_missing_items_json"]


def test_analytics_covers_product_batch_grade_responsibility_and_cycle_time(db):
    user = _user(db, "api-analytics")
    with _client(db, user, ["aftersales:read", "aftersales:write"]) as client:
        first_id = client.post("/api/aftersales/cases", json=_payload()).json()["data"]["id"]
        second_id = client.post(
            "/api/aftersales/cases",
            json={
                **_payload(),
                "customer_grade": "B",
                "order_id": "ORD002",
                "order_no_snapshot": "NO002",
                "batch_no": "BATCH-2",
            },
        ).json()["data"]["id"]
    first = db.get(AfterSalesCase, first_id)
    second = db.get(AfterSalesCase, second_id)
    first.responsibility_class = "A"
    first.has_compensation = True
    first.estimated_compensation_usd = Decimal("120")
    first.approved_at = first.created_at + timedelta(hours=6)
    first.closed_at = first.created_at + timedelta(hours=30)
    first.current_status = "closed"
    second.responsibility_class = "D"
    db.commit()

    # 2026-07-12 售后分析页独立码：分析端点改为 aftersales_analytics:read 把守
    with _client(db, user, ["aftersales_analytics:read", "aftersales:read_all"]) as client:
        response = client.get("/api/aftersales/analytics/summary")

    data = response.json()["data"]
    assert response.status_code == 200
    assert data["total"] == 2
    assert data["compensation_case_count"] == 1
    assert data["average_resolution_hours"] == 30.0
    assert {item["name"] for item in data["by_customer_grade"]} == {"A", "B"}
    assert {item["name"] for item in data["by_responsibility"]} == {"A", "D"}
    assert {item["name"] for item in data["by_batch"]} == {"BATCH-1", "BATCH-2"}
